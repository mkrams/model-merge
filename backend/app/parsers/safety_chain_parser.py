"""Parse imported files and extract safety chain items (Hazard → TestCase).

Supports: ReqIF, SysML v2, CSV/TSV, Excel (.xlsx/.xls), Word (.docx).
Strategy: heuristic classification based on naming patterns, attribute names,
package names, and requirement ID prefixes.
"""
from __future__ import annotations
import logging
import re
import csv
import io
import tempfile
from pathlib import Path
from typing import Optional

from ..models.safety import (
    SafetyProject, SafetyChain, Hazard, HazardousEvent,
    ASILDetermination, SafetyGoal, FSR, TestCase, FailureMode,
    compute_asil,
)
from ..parsers.sysml_v2_parser import parse_sysml_v2
from ..parsers.reqif_parser import parse_reqif

logger = logging.getLogger(__name__)

# ── Classification heuristics ────────────────────────────────────

_HAZ_PATTERNS = re.compile(r'hazard|haz[\-_]|malfunct|failure\s*mode', re.I)
_HE_PATTERNS = re.compile(r'hazardous\s*event|haz[\-_]?event|operating\s*situation', re.I)
_SG_PATTERNS = re.compile(r'safety\s*goal|safe[\-_]?goal|sg[\-_]', re.I)
_FSR_PATTERNS = re.compile(r'functional\s*safety|fsr[\-_]|safety\s*req|safe[\-_]?req', re.I)
_TC_PATTERNS = re.compile(r'test[\-_\s]?case|tc[\-_]|verification|val[\-_]?test', re.I)
_ASIL_PATTERN = re.compile(r'ASIL[\-_\s]*(QM|A|B|C|D)', re.I)
_FM_PATTERNS = re.compile(r'failure\s*mode|fmea|fm[\-_]', re.I)


def _classify_requirement(name: str, doc: str, req_id: str, pkg_name: str) -> str:
    """Classify a requirement into a chain level based on heuristics.
    Returns: 'hazard', 'hazardous_event', 'safety_goal', 'fsr', 'test_case', or 'unknown'.
    """
    text = f"{name} {doc} {req_id} {pkg_name}".lower()

    # Check ID prefix first (most reliable)
    rid = (req_id or "").upper()
    if rid.startswith("HAZ-") or rid.startswith("H-"):
        return "hazard"
    if rid.startswith("HE-"):
        return "hazardous_event"
    if rid.startswith("SG-"):
        return "safety_goal"
    if rid.startswith("FSR-") or rid.startswith("SR-"):
        return "fsr"
    if rid.startswith("TC-") or rid.startswith("VT-"):
        return "test_case"
    if rid.startswith("FM-"):
        return "failure_mode"

    # Then check name/doc patterns
    if _TC_PATTERNS.search(text):
        return "test_case"
    if _FSR_PATTERNS.search(text):
        return "fsr"
    if _SG_PATTERNS.search(text):
        return "safety_goal"
    if _HE_PATTERNS.search(text):
        return "hazardous_event"
    if _HAZ_PATTERNS.search(text):
        return "hazard"
    if _FM_PATTERNS.search(text):
        return "failure_mode"

    # Default: treat as FSR (most common requirement type)
    return "fsr"


def _extract_asil(text: str) -> Optional[str]:
    """Try to extract ASIL level from text."""
    m = _ASIL_PATTERN.search(text)
    if m:
        return m.group(1).upper()
    return None


# ── SysML v2 Import ─────────────────────────────────────────────

def parse_sysml_safety(raw_text: str, filename: str) -> SafetyProject:
    """Parse SysML v2 file and extract safety chain items."""
    model = parse_sysml_v2(raw_text, filename)
    project = SafetyProject(name=filename, source_filename=filename)

    # Collect all requirements and connections
    requirements = []
    connections = []
    _collect_from_packages(model.packages, requirements, connections, "")

    # Classify requirements
    classified: dict[str, list] = {
        "hazard": [], "hazardous_event": [], "safety_goal": [],
        "fsr": [], "test_case": [], "failure_mode": [], "unknown": [],
    }
    for req in requirements:
        level = _classify_requirement(
            req["name"], req.get("doc", ""), req.get("req_id", ""), req.get("package", "")
        )
        classified[level].append(req)

    # Build connection index (source → targets by kind)
    link_map: dict[str, list[tuple[str, str]]] = {}  # name → [(target, kind)]
    for conn in connections:
        src = conn.get("source", "")
        tgt = conn.get("target", "")
        kind = conn.get("kind", "")
        if src:
            link_map.setdefault(src, []).append((tgt, kind))

    # Build chains: start from hazards, follow links downstream
    for haz_data in classified["hazard"]:
        chain = SafetyChain()
        chain.hazard = Hazard(
            name=haz_data["name"],
            description=haz_data.get("doc", ""),
            status="draft",
        )
        asil_text = haz_data.get("doc", "") + " " + haz_data["name"]
        asil_level = _extract_asil(asil_text)
        if asil_level:
            chain.asil_determination = ASILDetermination(asil_level=asil_level)

        # Try to find downstream items via links
        _try_link_chain(chain, haz_data["name"], classified, link_map)
        project.chains.append(chain)

    # Any safety goals not linked to a hazard → create partial chains
    linked_sg_names = {c.safety_goal.name for c in project.chains if c.safety_goal}
    for sg_data in classified["safety_goal"]:
        if sg_data["name"] not in linked_sg_names:
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")  # Gap — no hazard
            chain.safety_goal = SafetyGoal(
                name=sg_data["name"],
                description=sg_data.get("doc", ""),
                asil_level=_extract_asil(sg_data.get("doc", "")) or "",
                status="draft",
            )
            _try_link_chain_downstream(chain, sg_data["name"], classified, link_map)
            project.chains.append(chain)

    # Any FSRs not in any chain → create minimal chains
    linked_fsr_names = {c.fsr.name for c in project.chains if c.fsr}
    for fsr_data in classified["fsr"]:
        if fsr_data["name"] not in linked_fsr_names:
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")
            chain.safety_goal = SafetyGoal(status="gap")
            chain.fsr = FSR(
                name=fsr_data["name"],
                description=fsr_data.get("doc", ""),
                asil_level=_extract_asil(fsr_data.get("doc", "")) or "",
                status="draft",
            )
            project.chains.append(chain)

    # Failure modes
    for fm_data in classified["failure_mode"]:
        project.failure_modes.append(FailureMode(
            name=fm_data["name"],
            description=fm_data.get("doc", ""),
        ))

    # Fill gaps on all chains
    _fill_chain_gaps(project)

    logger.info(f"Parsed SysML safety: {len(project.chains)} chains, {project.total_gaps} gaps")
    return project


def _collect_from_packages(packages, requirements, connections, parent_pkg):
    """Recursively collect requirements and connections from packages."""
    for pkg in packages:
        pkg_name = pkg.name if hasattr(pkg, 'name') else str(pkg)
        for req in getattr(pkg, 'requirement_defs', []):
            doc = getattr(req, 'doc', '') or ''
            raw = getattr(req, 'raw', '') or ''
            # Fallback: build description from attributes if doc is empty
            attrs = [(a.name, a.default_value or '') for a in getattr(req, 'attributes', [])]
            if not doc:
                # Try common attribute names for the requirement text
                for aname, aval in attrs:
                    aname_l = aname.lower()
                    if aval and any(k in aname_l for k in ['description', 'text', 'content', 'reqif.text', 'object text']):
                        doc = aval
                        break
            # Still empty? Use raw text or first attribute with substantial value
            if not doc and raw:
                doc = raw[:500]
            if not doc:
                for aname, aval in attrs:
                    if aval and len(aval) > 10:
                        doc = aval
                        break
            requirements.append({
                "name": req.name,
                "doc": doc,
                "req_id": getattr(req, 'req_id', ''),
                "package": pkg_name,
                "constraints": [c.expression for c in getattr(req, 'constraints', [])],
                "attributes": attrs,
                "raw": raw,
            })
        for conn in getattr(pkg, 'connections', []):
            connections.append({
                "source": getattr(conn, 'source', ''),
                "target": getattr(conn, 'target', ''),
                "kind": getattr(conn, 'kind', ''),
            })
        _collect_from_packages(getattr(pkg, 'subpackages', []), requirements, connections, pkg_name)


def _try_link_chain(chain, hazard_name, classified, link_map):
    """Try to build chain by following satisfy/verify/derive links from a hazard."""
    targets = link_map.get(hazard_name, [])
    for tgt, kind in targets:
        # Find what this target is
        for sg in classified["safety_goal"]:
            if sg["name"] == tgt and not chain.safety_goal:
                chain.safety_goal = SafetyGoal(
                    name=sg["name"], description=sg.get("doc", ""),
                    hazard_id=chain.hazard.id if chain.hazard else "",
                    asil_level=_extract_asil(sg.get("doc", "")) or "",
                    status="draft",
                )
                _try_link_chain_downstream(chain, tgt, classified, link_map)
                break


def _try_link_chain_downstream(chain, from_name, classified, link_map):
    """Follow links downstream from safety goal to FSR to test case."""
    targets = link_map.get(from_name, [])
    for tgt, kind in targets:
        for fsr in classified["fsr"]:
            if fsr["name"] == tgt and not chain.fsr:
                chain.fsr = FSR(
                    name=fsr["name"], description=fsr.get("doc", ""),
                    safety_goal_id=chain.safety_goal.id if chain.safety_goal else "",
                    status="draft",
                )
                # Look for test cases linked to this FSR
                for tc_tgt, tc_kind in link_map.get(tgt, []):
                    for tc in classified["test_case"]:
                        if tc["name"] == tc_tgt and not chain.test_case:
                            chain.test_case = TestCase(
                                name=tc["name"], description=tc.get("doc", ""),
                                fsr_id=chain.fsr.id,
                                status="draft",
                            )
                            break
                break


# ── ReqIF Import ─────────────────────────────────────────────────

def parse_reqif_safety(raw_text: str, filename: str) -> SafetyProject:
    """Parse ReqIF file and extract safety chain items."""
    model = parse_reqif(raw_text, filename)
    project = SafetyProject(name=filename, source_filename=filename)

    requirements = []
    connections = []
    _collect_from_packages(model.packages, requirements, connections, "")

    # Same classification logic as SysML
    classified: dict[str, list] = {
        "hazard": [], "hazardous_event": [], "safety_goal": [],
        "fsr": [], "test_case": [], "failure_mode": [], "unknown": [],
    }
    for req in requirements:
        level = _classify_requirement(
            req["name"], req.get("doc", ""), req.get("req_id", ""), req.get("package", "")
        )
        classified[level].append(req)

    link_map: dict[str, list[tuple[str, str]]] = {}
    for conn in connections:
        src = conn.get("source", "")
        tgt = conn.get("target", "")
        kind = conn.get("kind", "")
        if src:
            link_map.setdefault(src, []).append((tgt, kind))

    # Build chains from hazards
    for haz_data in classified["hazard"]:
        chain = SafetyChain()
        chain.hazard = Hazard(name=haz_data["name"], description=haz_data.get("doc", ""), status="draft")
        asil_level = _extract_asil(haz_data.get("doc", "") + " " + haz_data["name"])
        if asil_level:
            chain.asil_determination = ASILDetermination(asil_level=asil_level)
        _try_link_chain(chain, haz_data["name"], classified, link_map)
        project.chains.append(chain)

    # Unlinked safety goals
    linked_sg = {c.safety_goal.name for c in project.chains if c.safety_goal}
    for sg in classified["safety_goal"]:
        if sg["name"] not in linked_sg:
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")
            chain.safety_goal = SafetyGoal(name=sg["name"], description=sg.get("doc", ""), status="draft")
            project.chains.append(chain)

    # Unlinked FSRs
    linked_fsr = {c.fsr.name for c in project.chains if c.fsr}
    for fsr in classified["fsr"]:
        if fsr["name"] not in linked_fsr:
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")
            chain.safety_goal = SafetyGoal(status="gap")
            chain.fsr = FSR(name=fsr["name"], description=fsr.get("doc", ""), status="draft")
            project.chains.append(chain)

    # If no hazards found, create chains from whatever we have
    if not classified["hazard"] and not classified["safety_goal"]:
        for fsr in classified["fsr"]:
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")
            chain.safety_goal = SafetyGoal(status="gap")
            chain.fsr = FSR(name=fsr["name"], description=fsr.get("doc", ""), status="draft")
            project.chains.append(chain)

    # If nothing classified at all, make chains from "unknown" requirements
    if not project.chains:
        for req in classified.get("unknown", []) + classified.get("fsr", []):
            chain = SafetyChain()
            chain.hazard = Hazard(status="gap")
            chain.safety_goal = SafetyGoal(status="gap")
            chain.fsr = FSR(name=req["name"], description=req.get("doc", ""), status="draft")
            project.chains.append(chain)

    _fill_chain_gaps(project)
    logger.info(f"Parsed ReqIF safety: {len(project.chains)} chains, {project.total_gaps} gaps")
    return project


# ── CSV / Excel Import ───────────────────────────────────────────

def parse_csv_safety(raw_text: str, filename: str) -> SafetyProject:
    """Parse CSV/TSV and build safety chains.

    Supports TWO formats:
    1. Row-per-chain: columns like 'hazard', 'safety goal', 'fsr', 'test case'
    2. Row-per-item: columns like 'ID', 'Type', 'Name', 'Description', 'Parent_ID'
       Items are linked via Parent_ID to build chains.
    """
    project = SafetyProject(name=filename, source_filename=filename)
    reader = csv.DictReader(io.StringIO(raw_text))
    headers_raw = reader.fieldnames or []
    headers = [h.lower().strip() for h in headers_raw]

    # Detect format: does it have 'type' and 'id' columns?
    has_type = any("type" in h for h in headers)
    has_id = any(h in ("id", "req_id", "item_id") for h in headers)

    if has_type and has_id:
        return _parse_csv_row_per_item(reader, headers_raw, filename)

    # ── Format 1: Row-per-chain ──
    col_map = {}
    for h in headers:
        if any(k in h for k in ["hazard"]):
            col_map["hazard"] = h
        elif any(k in h for k in ["hazardous event", "haz event", "he "]):
            col_map["hazardous_event"] = h
        elif any(k in h for k in ["asil"]):
            col_map["asil"] = h
        elif any(k in h for k in ["safety goal", "sg"]):
            col_map["safety_goal"] = h
        elif any(k in h for k in ["fsr", "functional safety", "safety req"]):
            col_map["fsr"] = h
        elif any(k in h for k in ["test", "tc", "verification"]):
            col_map["test_case"] = h

    for row in reader:
        chain = SafetyChain()

        haz_text = row.get(col_map.get("hazard", ""), "").strip()
        if haz_text:
            chain.hazard = Hazard(name=haz_text, description=haz_text, status="draft")
        else:
            chain.hazard = Hazard(status="gap")

        he_text = row.get(col_map.get("hazardous_event", ""), "").strip()
        if he_text:
            chain.hazardous_event = HazardousEvent(
                name=he_text, description=he_text,
                hazard_id=chain.hazard.id if chain.hazard else "", status="draft",
            )

        asil_text = row.get(col_map.get("asil", ""), "").strip().upper()
        if asil_text and asil_text in ("QM", "A", "B", "C", "D"):
            chain.asil_determination = ASILDetermination(asil_level=asil_text)

        sg_text = row.get(col_map.get("safety_goal", ""), "").strip()
        if sg_text:
            chain.safety_goal = SafetyGoal(
                name=sg_text, description=sg_text,
                hazard_id=chain.hazard.id if chain.hazard else "",
                asil_level=asil_text or "", status="draft",
            )
        else:
            chain.safety_goal = SafetyGoal(status="gap")

        fsr_text = row.get(col_map.get("fsr", ""), "").strip()
        if fsr_text:
            chain.fsr = FSR(
                name=fsr_text, description=fsr_text,
                safety_goal_id=chain.safety_goal.id if chain.safety_goal else "",
                status="draft",
            )

        tc_text = row.get(col_map.get("test_case", ""), "").strip()
        if tc_text:
            chain.test_case = TestCase(
                name=tc_text, description=tc_text,
                fsr_id=chain.fsr.id if chain.fsr else "", status="draft",
            )

        if chain.hazard or chain.safety_goal or chain.fsr:
            project.chains.append(chain)

    _fill_chain_gaps(project)
    logger.info(f"Parsed CSV safety (row-per-chain): {len(project.chains)} chains, {project.total_gaps} gaps")
    return project


def _parse_csv_row_per_item(reader, headers_raw: list[str], filename: str) -> SafetyProject:
    """Parse CSV where each row is a single safety item with ID, Type, Name, Description, Parent_ID."""
    project = SafetyProject(name=filename, source_filename=filename)

    # Build column accessor (case-insensitive)
    def _col(row: dict, *candidates: str) -> str:
        for c in candidates:
            for key in row:
                if key.lower().strip() == c.lower():
                    return (row[key] or "").strip()
        return ""

    # First pass: collect all items by ID
    items: dict[str, dict] = {}
    rows_list = list(reader)
    for row in rows_list:
        item_id = _col(row, "id", "req_id", "item_id")
        if not item_id:
            continue
        item_type = _col(row, "type").lower()
        name = _col(row, "name", "title")
        desc = _col(row, "description", "text", "doc")
        asil = _col(row, "asil", "asil_level").upper()
        status = _col(row, "status").lower() or "draft"
        parent = _col(row, "parent_id", "parent", "derived_from")
        verified_by = _col(row, "verified_by", "verification")
        satisfied_by = _col(row, "satisfied_by", "satisfaction")

        # Classify type
        level = ""
        if any(k in item_type for k in ["hazardous event", "haz event", "he"]) or item_id.upper().startswith("HE-"):
            level = "hazardous_event"
        elif any(k in item_type for k in ["hazard"]) or item_id.upper().startswith("HAZ-"):
            level = "hazard"
        elif any(k in item_type for k in ["safety goal"]) or item_id.upper().startswith("SG-"):
            level = "safety_goal"
        elif any(k in item_type for k in ["fsr", "functional safety"]) or item_id.upper().startswith("FSR-"):
            level = "fsr"
        elif any(k in item_type for k in ["test", "tc"]) or item_id.upper().startswith("TC-"):
            level = "test_case"

        if not level:
            continue

        items[item_id] = {
            "id": item_id, "level": level, "name": name, "description": desc,
            "asil": asil if asil in ("QM", "A", "B", "C", "D") else "",
            "status": status if status in ("draft", "approved", "review", "gap") else "draft",
            "parent": parent, "verified_by": verified_by, "satisfied_by": satisfied_by,
        }

    # Second pass: build chains by tracing from hazards down
    hazards = {k: v for k, v in items.items() if v["level"] == "hazard"}
    used_ids: set[str] = set()

    for haz_id, haz in hazards.items():
        # Find hazardous events linked to this hazard
        hes = [v for v in items.values() if v["level"] == "hazardous_event" and v["parent"] == haz_id]
        if not hes:
            hes = [None]  # type: ignore

        for he in hes:
            # Find safety goals linked to this HE (or hazard)
            parent_for_sg = he["id"] if he else haz_id
            sgs = [v for v in items.values() if v["level"] == "safety_goal" and v["parent"] == parent_for_sg]
            if not sgs:
                sgs = [v for v in items.values() if v["level"] == "safety_goal" and v["parent"] == haz_id]
            if not sgs:
                sgs = [None]  # type: ignore

            for sg in sgs:
                parent_for_fsr = sg["id"] if sg else parent_for_sg
                fsrs = [v for v in items.values() if v["level"] == "fsr" and v["parent"] == parent_for_fsr]
                if not fsrs:
                    fsrs = [None]  # type: ignore

                for fsr_item in fsrs:
                    # Find test cases
                    tc = None
                    if fsr_item:
                        # Check Verified_By on the FSR
                        vb = fsr_item.get("verified_by", "")
                        if vb and vb in items:
                            tc = items[vb]
                        else:
                            # Check TCs that reference this FSR
                            tcs = [v for v in items.values() if v["level"] == "test_case" and v["parent"] == fsr_item["id"]]
                            tc = tcs[0] if tcs else None

                    chain = SafetyChain()

                    # Hazard
                    chain.hazard = Hazard(
                        name=haz["name"], description=haz["description"],
                        status=haz["status"], approved=(haz["status"] == "approved"),
                    )
                    used_ids.add(haz_id)

                    # HE
                    if he:
                        chain.hazardous_event = HazardousEvent(
                            name=he["name"], description=he["description"],
                            hazard_id=chain.hazard.id, status=he["status"],
                            approved=(he["status"] == "approved"),
                        )
                        used_ids.add(he["id"])

                    # ASIL (from highest priority source)
                    asil_val = ""
                    for src in [sg, he, haz]:
                        if src and src.get("asil"):
                            asil_val = src["asil"]
                            break
                    if asil_val:
                        chain.asil_determination = ASILDetermination(asil_level=asil_val, approved=True)

                    # SG
                    if sg:
                        chain.safety_goal = SafetyGoal(
                            name=sg["name"], description=sg["description"],
                            hazard_id=chain.hazard.id, asil_level=asil_val,
                            status=sg["status"], approved=(sg["status"] == "approved"),
                        )
                        used_ids.add(sg["id"])

                    # FSR
                    if fsr_item:
                        chain.fsr = FSR(
                            name=fsr_item["name"], description=fsr_item["description"],
                            safety_goal_id=chain.safety_goal.id if chain.safety_goal else "",
                            asil_level=asil_val, status=fsr_item["status"],
                            approved=(fsr_item["status"] == "approved"),
                        )
                        used_ids.add(fsr_item["id"])

                    # TC
                    if tc:
                        chain.test_case = TestCase(
                            name=tc["name"], description=tc["description"],
                            fsr_id=chain.fsr.id if chain.fsr else "",
                            status=tc["status"], approved=(tc["status"] == "approved"),
                        )
                        used_ids.add(tc["id"])

                    project.chains.append(chain)

    # Third pass: orphan items not in any chain — create partial chains for them
    for item_id, item in items.items():
        if item_id in used_ids:
            continue
        chain = SafetyChain()
        lvl = item["level"]
        asil_val = item.get("asil", "")

        if lvl == "hazard":
            chain.hazard = Hazard(name=item["name"], description=item["description"], status=item["status"])
        elif lvl == "hazardous_event":
            chain.hazardous_event = HazardousEvent(name=item["name"], description=item["description"], status=item["status"])
        elif lvl == "safety_goal":
            chain.safety_goal = SafetyGoal(name=item["name"], description=item["description"], asil_level=asil_val, status=item["status"])
        elif lvl == "fsr":
            chain.fsr = FSR(name=item["name"], description=item["description"], asil_level=asil_val, status=item["status"])
        elif lvl == "test_case":
            chain.test_case = TestCase(name=item["name"], description=item["description"], status=item["status"])

        if asil_val:
            chain.asil_determination = ASILDetermination(asil_level=asil_val)

        project.chains.append(chain)

    _fill_chain_gaps(project)
    logger.info(f"Parsed CSV safety (row-per-item): {len(project.chains)} chains, {project.total_gaps} gaps")
    return project


# ── Gap Filling ──────────────────────────────────────────────────

def _fill_chain_gaps(project: SafetyProject):
    """Ensure every chain has all 6 slots (even if gap)."""
    for chain in project.chains:
        if not chain.hazard:
            chain.hazard = Hazard(status="gap")
        if not chain.hazardous_event:
            chain.hazardous_event = HazardousEvent(
                hazard_id=chain.hazard.id, status="gap",
            )
        if not chain.asil_determination:
            chain.asil_determination = ASILDetermination()
        if not chain.safety_goal:
            chain.safety_goal = SafetyGoal(
                hazard_id=chain.hazard.id, status="gap",
            )
        if not chain.fsr:
            chain.fsr = FSR(
                safety_goal_id=chain.safety_goal.id, status="gap",
            )
        if not chain.test_case:
            chain.test_case = TestCase(
                fsr_id=chain.fsr.id, status="gap",
            )


# ── Excel Import ──────────────────────────────────────────────────

def parse_excel_safety(file_bytes: bytes, filename: str) -> SafetyProject:
    """Parse .xlsx/.xls file and build safety chains.

    Reads the first sheet (or a sheet named 'requirements', 'hazards', 'safety', etc.).
    Then converts to CSV-like rows and delegates to CSV parser logic.
    """
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl required for Excel import — pip install openpyxl")

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)

    # Pick the best sheet
    sheet = None
    preferred = ["requirements", "hazards", "safety", "asil", "fsr", "chains"]
    for name in wb.sheetnames:
        if any(p in name.lower() for p in preferred):
            sheet = wb[name]
            break
    if sheet is None:
        sheet = wb.active or wb[wb.sheetnames[0]]

    # Read all rows into CSV text
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return SafetyProject(name=filename, source_filename=filename)

    # First non-empty row as header
    header_row = None
    data_start = 0
    for i, row in enumerate(rows):
        if any(cell is not None for cell in row):
            header_row = row
            data_start = i + 1
            break

    if header_row is None:
        return SafetyProject(name=filename, source_filename=filename)

    # Build CSV string
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([str(h or "").strip() for h in header_row])
    for row in rows[data_start:]:
        if any(cell is not None for cell in row):
            writer.writerow([str(cell or "").strip() for cell in row])

    csv_text = output.getvalue()
    wb.close()

    # Parse using CSV logic
    project = parse_csv_safety(csv_text, filename)
    logger.info(f"Parsed Excel safety: {len(project.chains)} chains")
    return project


# ── Word (.docx) Import ──────────────────────────────────────────

def parse_docx_safety(file_bytes: bytes, filename: str) -> SafetyProject:
    """Parse .docx file and extract safety items from tables and text.

    Looks for:
    1. Tables with columns matching safety chain levels
    2. Structured headings like "Hazard:", "Safety Goal:", etc.
    3. Numbered/bulleted requirements with ID prefixes (HAZ-, SG-, FSR-, TC-)
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx required for Word import — pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))

    # Strategy 1: Try tables first
    for table in doc.tables:
        result = _parse_docx_table(table, filename)
        if result and len(result.chains) > 0:
            return result

    # Strategy 2: Parse structured text
    return _parse_docx_text(doc, filename)


def _parse_docx_table(table, filename: str) -> Optional[SafetyProject]:
    """Parse a Word table into safety items."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    if len(rows) < 2:
        return None

    # Use first row as headers
    headers = rows[0]
    if not any(h.lower() for h in headers if h):
        return None

    # Build CSV text from table
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    for row in rows[1:]:
        # Pad/trim to header length
        padded = row[:len(headers)] + [""] * max(0, len(headers) - len(row))
        writer.writerow(padded)

    csv_text = output.getvalue()

    try:
        project = parse_csv_safety(csv_text, filename)
        return project if project.chains else None
    except Exception:
        return None


_DOCX_LEVEL_PATTERNS = {
    "hazard": re.compile(r"(?:HAZ[-_]?\d+|Hazard\s*:)", re.IGNORECASE),
    "hazardous_event": re.compile(r"(?:HE[-_]?\d+|Hazardous\s+Event\s*:)", re.IGNORECASE),
    "safety_goal": re.compile(r"(?:SG[-_]?\d+|Safety\s+Goal\s*:)", re.IGNORECASE),
    "fsr": re.compile(r"(?:FSR[-_]?\d+|Functional\s+Safety\s+Req)", re.IGNORECASE),
    "test_case": re.compile(r"(?:TC[-_]?\d+|Test\s+Case\s*:)", re.IGNORECASE),
}


def _parse_docx_text(doc, filename: str) -> SafetyProject:
    """Parse structured text from Word document paragraphs."""
    project = SafetyProject(name=filename, source_filename=filename)

    items: list[dict] = []
    current_item: Optional[dict] = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this line starts a new item
        matched_level = None
        for level, pattern in _DOCX_LEVEL_PATTERNS.items():
            if pattern.search(text):
                matched_level = level
                break

        if matched_level:
            if current_item:
                items.append(current_item)

            # Extract ID if present
            id_match = re.match(r"((?:HAZ|HE|SG|FSR|TC)[-_]?\d+)\s*[:\-–]\s*(.*)", text, re.IGNORECASE)
            if id_match:
                item_id = id_match.group(1).upper().replace("_", "-")
                name = id_match.group(2).strip()
            else:
                item_id = f"{matched_level.upper()}-AUTO-{len(items)+1}"
                # Strip "Hazard:" prefix etc.
                name = re.sub(r"^(?:Hazard|Hazardous\s+Event|Safety\s+Goal|FSR|Test\s+Case)\s*:\s*",
                             "", text, flags=re.IGNORECASE).strip()

            current_item = {
                "id": item_id, "level": matched_level,
                "name": name, "description": "", "asil": "", "status": "draft",
                "parent": "",
            }
        elif current_item:
            # Continuation of current item — append as description
            if current_item["description"]:
                current_item["description"] += "\n" + text
            else:
                current_item["description"] = text

            # Check for ASIL mention
            asil_match = re.search(r"ASIL\s*[:\-]?\s*([ABCD]|QM)", text, re.IGNORECASE)
            if asil_match:
                current_item["asil"] = asil_match.group(1).upper()

    if current_item:
        items.append(current_item)

    if not items:
        return project

    # Build chains from items (simple grouping — try to link by proximity and ASIL)
    # Group by level and build chains
    by_level: dict[str, list[dict]] = {"hazard": [], "hazardous_event": [],
                                        "safety_goal": [], "fsr": [], "test_case": []}
    for item in items:
        if item["level"] in by_level:
            by_level[item["level"]].append(item)

    # Build chains: pair items by order (1st hazard → 1st SG → 1st FSR → 1st TC)
    max_len = max(len(v) for v in by_level.values()) if by_level else 0
    for i in range(max_len):
        chain = SafetyChain()

        haz = by_level["hazard"][i] if i < len(by_level["hazard"]) else None
        he = by_level["hazardous_event"][i] if i < len(by_level["hazardous_event"]) else None
        sg = by_level["safety_goal"][i] if i < len(by_level["safety_goal"]) else None
        fsr_item = by_level["fsr"][i] if i < len(by_level["fsr"]) else None
        tc = by_level["test_case"][i] if i < len(by_level["test_case"]) else None

        asil_val = ""
        for src in [sg, he, haz, fsr_item]:
            if src and src.get("asil"):
                asil_val = src["asil"]
                break

        if haz:
            chain.hazard = Hazard(name=haz["name"], description=haz.get("description", ""), status="draft")
        if he:
            chain.hazardous_event = HazardousEvent(name=he["name"], description=he.get("description", ""), status="draft")
        if asil_val:
            chain.asil_determination = ASILDetermination(asil_level=asil_val)
        if sg:
            chain.safety_goal = SafetyGoal(name=sg["name"], description=sg.get("description", ""), asil_level=asil_val, status="draft")
        if fsr_item:
            chain.fsr = FSR(name=fsr_item["name"], description=fsr_item.get("description", ""), status="draft")
        if tc:
            chain.test_case = TestCase(name=tc["name"], description=tc.get("description", ""), status="draft")

        project.chains.append(chain)

    _fill_chain_gaps(project)
    logger.info(f"Parsed DOCX safety: {len(project.chains)} chains, {project.total_gaps} gaps")
    return project


# ── Main entry point ─────────────────────────────────────────────

def parse_safety_chain(raw_text: str, filename: str) -> SafetyProject:
    """Auto-detect format and parse safety chain items.

    For text-based formats (CSV, ReqIF, SysML), raw_text is the file content.
    For binary formats (xlsx, docx), raw_text should be the base64/raw string
    — but those are handled via parse_safety_chain_bytes() instead.
    """
    fname = filename.lower()
    if fname.endswith(".reqif") or fname.endswith(".xml"):
        return parse_reqif_safety(raw_text, filename)
    elif fname.endswith(".csv") or fname.endswith(".tsv"):
        return parse_csv_safety(raw_text, filename)
    else:
        # Default: SysML v2
        return parse_sysml_safety(raw_text, filename)


def parse_safety_chain_bytes(file_bytes: bytes, filename: str) -> SafetyProject:
    """Parse binary file formats (Excel, Word) or fall back to text parser."""
    fname = filename.lower()
    if fname.endswith(".xlsx") or fname.endswith(".xls"):
        return parse_excel_safety(file_bytes, filename)
    elif fname.endswith(".docx"):
        return parse_docx_safety(file_bytes, filename)
    else:
        # Text-based: decode and use text parser
        raw_text = file_bytes.decode("utf-8", errors="replace")
        return parse_safety_chain(raw_text, filename)
